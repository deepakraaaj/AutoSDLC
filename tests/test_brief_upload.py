from io import BytesIO
from pathlib import Path
import sys

import pytest
from docx import Document as DocxDocument


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


from brief_upload import SUPPORTED_UPLOAD_EXTENSIONS, extract_uploaded_brief_text  # noqa: E402


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Project Aurora", level=1)
    document.add_paragraph("Build a simple dashboard for operations.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Team"
    table.cell(1, 1).text = "5"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_supported_upload_extensions_include_docx():
    assert ".md" in SUPPORTED_UPLOAD_EXTENSIONS
    assert ".docx" in SUPPORTED_UPLOAD_EXTENSIONS


def test_extract_uploaded_brief_text_from_docx():
    text = extract_uploaded_brief_text("brief.docx", _build_docx_bytes())

    assert "# Project Aurora" in text
    assert "Build a simple dashboard for operations." in text
    assert "Owner | Value" in text
    assert "Team | 5" in text


def test_extract_uploaded_brief_text_from_markdown():
    text = extract_uploaded_brief_text("brief.md", b"# Project: Demo\n\nSome content here.")

    assert text == "# Project: Demo\n\nSome content here."


def test_extract_uploaded_brief_text_rejects_unsupported_files():
    with pytest.raises(ValueError, match="Only \\.md and \\.docx files are accepted\\."):
        extract_uploaded_brief_text("brief.pdf", b"%PDF-1.4")
